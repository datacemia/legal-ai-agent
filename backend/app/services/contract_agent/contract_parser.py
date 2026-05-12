import os
import re
from typing import Optional

import fitz
import pdfplumber
from docx import Document

MAX_PAGES = 20
MAX_CHARS = 200_000


ARABIC_FIXES = {
    "لالستشارات": "للاستشارات",
    "األطراف": "الأطراف",
    "األضرار": "الأضرار",
    "اإلنهاء": "الإنهاء",
    "اإلخلال": "الإخلال",
    "اإلتفاق": "الاتفاق",
    "األول": "الأول",
    "األولى": "الأولى",
}


def clean_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\x00", " ")

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"[ \t]+\n", "\n", text)

    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove common page markers
    text = re.sub(r"(?im)^\s*page\s+\d+\s*$", "", text)
    text = re.sub(r"(?im)^\s*\d+\s*\|\s*page\s*$", "", text)
    text = re.sub(r"(?im)^\s*صفحة\s+\d+\s*$", "", text)

    # Remove long signature/underline artifacts
    text = re.sub(r"_{3,}", "", text)
    text = re.sub(r"-{5,}", "", text)

    for wrong, correct in ARABIC_FIXES.items():
        text = text.replace(wrong, correct)

    return text.strip()[:MAX_CHARS]


def text_quality_score(text: str) -> int:
    if not text or not text.strip():
        return 0

    score = 0
    stripped = text.strip()

    if len(stripped) > 500:
        score += 30
    elif len(stripped) > 100:
        score += 15

    if "\n" in stripped:
        score += 10

    if re.search(r"\b(agreement|contract|article|section|clause)\b", stripped, re.I):
        score += 20

    if re.search(r"(المادة|العقد|الاتفاق|البند|الطرف)", stripped):
        score += 20

    if re.search(r"\d+(\.\d+)?", stripped):
        score += 10

    # Penalize extraction noise
    weird_chars = len(re.findall(r"[�□■●◆]", stripped))
    if weird_chars > 10:
        score -= 20

    very_short_lines = [
        line for line in stripped.splitlines()
        if 0 < len(line.strip()) <= 2
    ]
    if len(very_short_lines) > 30:
        score -= 15

    return max(score, 0)


def extract_text_from_pdf_pymupdf(file_path: str) -> str:
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise ValueError(f"Failed to open PDF with PyMuPDF: {e}")

    text_parts = []

    try:
        total_pages = min(len(doc), MAX_PAGES)

        for page_num in range(total_pages):
            page = doc.load_page(page_num)
            page_text = page.get_text("text") or ""

            if page_text.strip():
                text_parts.append(page_text)

            if sum(len(part) for part in text_parts) > MAX_CHARS:
                break

    finally:
        doc.close()

    text = clean_text("\n".join(text_parts))

    if not text:
        raise ValueError("Empty or unreadable PDF with PyMuPDF")

    return text


def extract_text_from_pdf_pdfplumber(file_path: str) -> str:
    text_parts = []

    try:
        with pdfplumber.open(file_path) as pdf:
            total_pages = min(len(pdf.pages), MAX_PAGES)

            for i in range(total_pages):
                page = pdf.pages[i]

                page_text = page.extract_text(
                    x_tolerance=1,
                    y_tolerance=3,
                ) or ""

                if page_text.strip():
                    text_parts.append(page_text)

                if sum(len(part) for part in text_parts) > MAX_CHARS:
                    break

    except Exception as e:
        raise ValueError(f"Failed to process PDF with pdfplumber: {e}")

    text = clean_text("\n".join(text_parts))

    if not text:
        raise ValueError("Empty or unreadable PDF with pdfplumber")

    return text


def extract_text_from_pdf(file_path: str) -> str:
    candidates = []

    for extractor in (
        extract_text_from_pdf_pymupdf,
        extract_text_from_pdf_pdfplumber,
    ):
        try:
            extracted = extractor(file_path)
            candidates.append(extracted)
        except Exception:
            continue

    if not candidates:
        raise ValueError("Failed to extract readable text from PDF")

    best_text = max(candidates, key=text_quality_score)

    if text_quality_score(best_text) < 20:
        raise ValueError("PDF text quality is too low or unreadable")

    return best_text


def extract_docx_tables(doc: Document) -> list[str]:
    table_texts = []

    for table in doc.tables:
        rows = []

        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            ]

            if cells:
                rows.append(" | ".join(cells))

        if rows:
            table_texts.append("\n".join(rows))

    return table_texts


def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)

        paragraphs = [
            p.text.strip()
            for p in doc.paragraphs
            if p.text and p.text.strip()
        ]

        tables = extract_docx_tables(doc)

        text = "\n".join(paragraphs + tables)

    except Exception as e:
        raise ValueError(f"Failed to process DOCX: {e}")

    text = clean_text(text)

    if not text:
        raise ValueError("Empty or unreadable DOCX")

    return text


def extract_text(file_path: str, file_type: Optional[str]) -> str:
    if not file_path or not os.path.exists(file_path):
        raise ValueError("File does not exist")

    normalized_type = (file_type or "").lower().strip().replace(".", "")

    if normalized_type == "pdf":
        return extract_text_from_pdf(file_path)

    if normalized_type == "docx":
        return extract_text_from_docx(file_path)

    raise ValueError("Unsupported file type")