import os
import io
import re
from io import BytesIO

import fitz
import pytesseract
from PIL import Image
from docx import Document
from fastapi import UploadFile

from langdetect import detect
from openai import OpenAI


# =========================
# OPENAI CLIENT
# =========================

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# =========================
# TESSERACT CONFIG
# =========================

TESSERACT_CMD = os.getenv("TESSERACT_CMD")
OCR_LANGS = os.getenv("OCR_LANGS", "ara+fra+eng")
OCR_CONFIG = os.getenv(
    "OCR_CONFIG",
    "--oem 3 --psm 4 -c preserve_interword_spaces=1"
)

if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
elif os.name == "nt":
    pytesseract.pytesseract.tesseract_cmd = (
        r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    )


# =========================
# BASIC CLEANING (SAFE)
# =========================

def fix_arabic_word_spacing(text: str) -> str:
    text = re.sub(r"\bال\s+([اأإآء-ي])", r"ال\1", text)
    text = re.sub(r"\b([لبكفو])\s+([اأإآ])", r"\1\2", text)
    return text


def fix_latin_spacing(text: str) -> str:
    text = re.sub(r"([a-zà-ÿ])([A-Z])", r"\1 \2", text)
    text = re.sub(r"([A-Z]+)([A-Z][a-zà-ÿ])", r"\1 \2", text)
    return text


def remove_ocr_noise(text: str) -> str:
    arabic_chars = len(re.findall(r"[\u0600-\u06FF]", text))
    latin_chars = len(re.findall(r"[A-Za-z]", text))

    if arabic_chars > latin_chars:
        text = re.sub(r"\b[A-Za-z]{2,}\b", " ", text)

    return text


def normalize_ocr_text(text: str) -> str:
    replacements = {
        "\x00": " ",
        "\u200f": "",
        "\u200e": "",
        "\ufeff": "",
        "ﻻ": "لا",
        "ﻷ": "لأ",
        "ﻹ": "لإ",
        "ﻵ": "لآ",
        "ﻼ": "لا",
        "ـ": "",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    text = fix_latin_spacing(text)
    text = fix_arabic_word_spacing(text)
    text = remove_ocr_noise(text)

    text = re.sub(r"[ \t]+", " ", text)

    return text.strip()


def clean_extracted_text(text: str) -> str:
    text = normalize_ocr_text(text)

    lines = []
    for line in text.splitlines():
        line = line.strip()
        if line:
            lines.append(" ".join(line.split()))

    return "\n".join(lines)


# =========================
# LANGUAGE DETECTION
# =========================

def detect_language(text: str) -> str:
    try:
        return detect(text)
    except:
        return "unknown"


# =========================
# AI CLEANING (SMART)
# =========================

def ai_clean_text(text: str, lang: str) -> str:
    """
    Intelligent cleanup using LLM.
    Fix OCR mistakes WITHOUT changing meaning.
    """

    if len(text) < 200:
        return text  # skip small texts

    prompt = f"""
Clean this OCR text in {lang}.

Rules:
- Fix spacing issues
- Fix broken words
- Remove noise
- Keep original meaning
- DO NOT summarize
- DO NOT translate

Text:
{text[:4000]}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print("AI cleaning error:", e)
        return text


# =========================
# DOCX
# =========================

def extract_docx_text(content: bytes) -> str:
    try:
        doc = Document(BytesIO(content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return clean_extracted_text("\n".join(text_parts))

    except Exception as e:
        print("DOCX error:", e)
        return ""


# =========================
# OCR PDF
# =========================

def extract_scanned_pdf_text(content: bytes) -> str:
    pdf = fitz.open(stream=content, filetype="pdf")
    pages_text = []

    for page in pdf:
        pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
        img = Image.open(io.BytesIO(pix.tobytes("png")))

        text = pytesseract.image_to_string(
            img,
            lang=OCR_LANGS,
            config=OCR_CONFIG
        )

        text = clean_extracted_text(text)

        if text:
            pages_text.append(text)

    pdf.close()
    return "\n\n".join(pages_text)


# =========================
# PDF TEXT
# =========================

def extract_pdf_text(content: bytes) -> str:
    pdf = fitz.open(stream=content, filetype="pdf")
    pages_text = []

    for page in pdf:
        text = page.get_text("text", sort=True)
        text = clean_extracted_text(text)

        if text:
            pages_text.append(text)

    pdf.close()
    return "\n\n".join(pages_text)


# =========================
# MAIN ENTRY
# =========================

async def extract_study_text(file: UploadFile) -> str:
    content = await file.read()
    filename = (file.filename or "").lower()

    try:
        if filename.endswith(".pdf"):
            text = extract_pdf_text(content)

            if not text or len(text) < 80:
                text = extract_scanned_pdf_text(content)

        elif filename.endswith(".docx"):
            text = extract_docx_text(content)

        else:
            return ""

        if not text:
            return ""

        # 🔥 language detection
        lang = detect_language(text[:1000])
        print("Detected language:", lang)

        # 🔥 AI smart cleaning
        text = ai_clean_text(text, lang)

        return text.strip()

    except Exception as e:
        print("Error:", e)
        return ""