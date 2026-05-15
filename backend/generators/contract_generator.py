import random
from pathlib import Path
from faker import Faker
from docx import Document

fake_en = Faker("en_US")
fake_fr = Faker("fr_FR")

OUTPUT_DIR = Path("generated_contracts")
OUTPUT_DIR.mkdir(exist_ok=True)

DOMAINS = [
    "employment",
    "saas",
    "loan",
    "distribution",
    "commercial_lease",
    "nda",
    "consulting",
    "vendor",
    "licensing",
    "partnership",
]

LANGUAGES = ["en", "fr", "ar"]

EN_CLAUSES = {
    "payment": "The Client shall pay all invoices within 30 days.",
    "termination": "Either party may terminate this Agreement upon material breach.",
    "confidentiality": "Confidential information must not be disclosed.",
    "liability": "Liability shall not exceed the total fees paid.",
    "ip": "All intellectual property remains the property of the Provider.",
}

FR_CLAUSES = {
    "payment": "Le Client doit payer les factures dans un délai de 30 jours.",
    "termination": "Chaque partie peut résilier le contrat en cas de manquement grave.",
    "confidentiality": "Les informations confidentielles ne doivent pas être divulguées.",
    "liability": "La responsabilité est limitée aux montants payés.",
    "ip": "La propriété intellectuelle reste la propriété du Prestataire.",
}

AR_CLAUSES = {
    "payment": "يلتزم العميل بدفع الفواتير خلال 30 يوماً.",
    "termination": "يجوز لأي طرف إنهاء العقد في حالة الإخلال الجوهري.",
    "confidentiality": "يجب عدم الكشف عن المعلومات السرية.",
    "liability": "تقتصر المسؤولية على المبالغ المدفوعة.",
    "ip": "تبقى الملكية الفكرية ملكاً لمقدم الخدمة.",
}


def get_clauses(language):
    if language == "fr":
        return FR_CLAUSES

    if language == "ar":
        return AR_CLAUSES

    return EN_CLAUSES


def generate_title(domain, language):

    titles = {
        "en": {
            "employment": "Employment Agreement",
            "saas": "SaaS Services Agreement",
            "loan": "Loan Agreement",
            "distribution": "Distribution Agreement",
            "commercial_lease": "Commercial Lease Agreement",
        },

        "fr": {
            "employment": "Contrat de Travail",
            "saas": "Contrat SaaS",
            "loan": "Contrat de Prêt",
            "distribution": "Contrat de Distribution",
            "commercial_lease": "Bail Commercial",
        },

        "ar": {
            "employment": "عقد عمل",
            "saas": "عقد خدمات سحابية",
            "loan": "عقد قرض",
            "distribution": "عقد توزيع",
            "commercial_lease": "عقد كراء تجاري",
        }
    }

    return titles.get(
        language,
        titles["en"]
    ).get(domain, domain)


def generate_contract(index, language, domain):

    doc = Document()

    title = generate_title(domain, language)

    doc.add_heading(title, level=1)

    company_1 = fake_en.company()
    company_2 = fake_en.company()

    if language == "fr":

        intro = (
            f"Le présent contrat est conclu entre "
            f"{company_1} et {company_2}."
        )

    elif language == "ar":

        intro = (
            f"تم إبرام هذا العقد بين "
            f"{company_1} و {company_2}."
        )

    else:

        intro = (
            f"This Agreement is entered into between "
            f"{company_1} and {company_2}."
        )

    doc.add_paragraph(intro)

    clauses = get_clauses(language)

    selected = random.sample(
        list(clauses.items()),
        k=len(clauses)
    )

    for i, (clause_name, clause_text) in enumerate(
        selected,
        start=1
    ):

        if language == "fr":

            clause_title = (
                f"Article {i} - {clause_name}"
            )

        elif language == "ar":

            clause_title = (
                f"المادة {i} - {clause_name}"
            )

        else:

            clause_title = (
                f"Section {i} - "
                f"{clause_name.title()}"
            )

        doc.add_heading(
            clause_title,
            level=2
        )

        doc.add_paragraph(clause_text)

    filename = (
        OUTPUT_DIR /
        f"{language}_{domain}_{index}.docx"
    )

    doc.save(filename)

    print(f"Generated: {filename}")


TOTAL = 1000

for i in range(TOTAL):

    language = random.choice(LANGUAGES)
    domain = random.choice(DOMAINS)

    generate_contract(
        i + 1,
        language,
        domain
    )

print("DONE")